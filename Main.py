import docx
from docx import Document
import xlwt
import xlrd
import os
import openpyxl
from openpyxl import Workbook
from openpyxl.utils import column_index_from_string
import win32com
from win32com import client
from win32com.client import Dispatch, constants
import traceback
def findNum(paragraphs,Begin,Str):#从paragraphs的第Begin行开始找，找到匹配的
    if Begin<0:
        Begin=0;
    for i in range(Begin,len(paragraphs)):
        for j in range(len(paragraphs[i].text)):
            if paragraphs[i].text[j]==Str[0] and paragraphs[i].text[j:j+len(Str)]==Str:
                return i
    return -1

def tableFindNum(table,Begin,Str):#从table的第Begin行开始找，找到匹配的字符串。例：如果传入“葡萄园建立时间”.
    for i in range(Begin,len(table.rows)):
        for j in range(len(table.rows[i].cells)):
            if Str in table.rows[i].cells[j].text:
                return [i,j]
    return [-1,-1]

def Digit(Str):#整数化：如果传入字符串为整数，返回对应整数；否则返回原值
    if Str.isdigit():
        return int(Str)
    return Str

def Cost(Str):#
    if Str=='':return -1
    cb=['水','化肥','农药','套袋','电','有机肥','生长调节剂','塑料薄膜']
    for z in range(len(cb)):
        if Str in cb[z]:return z
    return -1

def num_Chinese(word):
    Sum=0
    for ch in word:
        if '\u4e00' <= ch <= '\u9fff':
            Sum=Sum+1
    return Sum
def is_AllChinese(word):
    if word=='':return True
    for ch in word:
        if '\u4e00' > ch or ch > '\u9fff':
             return False
    return True

def Extract(word_file,xlsx_file='',Start_line=0):#传入单个Word和模板的地址，将该Word的信息填入模板
    word_file=word_file.replace('\\','//')
    print(word_file)
    global CompletedNum,Dict,Template_xlsx,Name
    if xlsx_file=='':
        xlsx_file=Template_xlsx
    #print(word_file)
    wbook=openpyxl.load_workbook(xlsx_file)
    Name=wbook.sheetnames
    doc = Document(word_file)#用python-docx打开文件

    Demonstration=0 #Demonstration=1,则为示范园
    if '示范园' in word_file:
        if '非示范园' in word_file:
            Demonstration=0
        else:Demonstration=1
    else:
        if findNum(doc.paragraphs,Start_line+0,"非示范园")-Start_line<=3:
            Demonstration=0
        elif findNum(doc.paragraphs,Start_line+0,"示范园基本情况")-Start_line<=3:
            Demonstration=1
    
    
    for j in range(len(doc.paragraphs[0].text)):
        if doc.paragraphs[0].text[j]=='示':
            if doc.paragraphs[0].text[j-1]=='非':
                Demonstration=0
            else :
                Demonstration=1
    #if Demonstration==1:
    i=findNum(doc.paragraphs,Start_line+0,"所属试验站")
    temp=doc.paragraphs[i].text
    temp=temp[temp.find('：')+1:].replace(' ','')
    #print(1,temp)
    str1=''
    temp=temp.replace('国家葡萄产业技术体系','')
    temp=temp.replace('综合试验站','')
    if temp!='':
        str1='国家葡萄产业技术体系'+temp+'综合试验站'
    #print(str1)
    i=findNum(doc.paragraphs,i+1,"地址")
    temp=doc.paragraphs[i].text
    j=len(temp)
    if '，' in temp:
        j=temp.find('，')
    #if j==-1:
    str2=temp[temp.find('：')+1:j].replace(' ','')
    for Nsheet in range(len(Name)):
        if Name[Nsheet] in str2 or Name[Nsheet] in word_file:
            break;
    if Nsheet>=len(Name):Nsheet=len(Name)-1
    str3=Name[Nsheet];
    #if '北京' in str2 or '上海' in str2 or '重庆' in str2 or'天津' in str2:str2=str2[str2.find('省')+1:]
    if str2[0]=='省':str2=str2[1:]
    print(str2)
    
    CompletedNum[Nsheet]=CompletedNum[Nsheet]+1
    Row=CompletedNum[Nsheet]+1#第一行是表头，所以从第二行开始
    print(CompletedNum[Nsheet],Row,str3,str2)

    wsheet=wbook[str3]
    wsheet.cell(Row,column_index_from_string("FS")).value=word_file
    #wsheet=wbook.active
    #wsheet.cell(Row,B)#行，列
    #print(Row)
    wsheet.cell(Row,column_index_from_string("B")).value=str3
    wsheet.cell(Row,column_index_from_string("G")).value=str1
    wsheet.cell(Row,column_index_from_string("H")).value=str2
    
    str1=1
    if Demonstration !=1:
        str1=2
    wsheet.cell(Row,column_index_from_string("I")).value=str1
    i=findNum(doc.paragraphs,Start_line,"按设施的功能")
    a=[1]*4
    for j in range(len(doc.paragraphs[i].text)):
        if doc.paragraphs[i].text[j]=='□':
            str1=doc.paragraphs[i].text[j-4:j]
            if str1=='促早栽培':a[1]=0
            elif str1=='延迟栽培':a[2]=0
            elif str1=='避雨栽培':a[3]=0
    x=0
    for j in range(1,4):
        if a[j]==1:
            x=x*10+j
    if x!=0:wsheet.cell(Row,column_index_from_string("D")).value=x
    
    
    i=findNum(doc.paragraphs,Start_line,"设施的结构")
    b=[1]*5
    for j in range(len(doc.paragraphs[i].text)):
        if doc.paragraphs[i].text[j]=='□':
            if j<=len(doc.paragraphs[i].text)-2 and doc.paragraphs[i].text[j+1]=='√':continue;
            str1=doc.paragraphs[i].text[j-4:j]
            if str1=='塑料大棚':b[1]=0
            elif str1=='日光温室':b[2]=0
            elif str1=='加温温室':b[3]=0
            elif str1=='易避雨棚':b[4]=0
    y=0
    for j in range(1,5):
        if b[j]==1:
            y=y*10+j
    if y!=0: wsheet.cell(Row,column_index_from_string("E")).value=y
    if x+y>0:wsheet.cell(Row,column_index_from_string("C")).value=2
    else : wsheet.cell(Row,column_index_from_string("C")).value=1

    i=findNum(doc.paragraphs,i+1,"观光采摘葡萄园")
    a=[1]*3
    for j in range(len(doc.paragraphs[i].text)):
        if doc.paragraphs[i].text[j]=='□':
            if j<=len(doc.paragraphs[i].text)-2 and doc.paragraphs[i].text[j+1]=='√':continue;
            if doc.paragraphs[i].text[j-1]=='否':a[2]=0
            elif doc.paragraphs[i].text[j-1]=='是':a[1]=0
    if a[1]==1: wsheet.cell(Row,column_index_from_string("F")).value=1
    if a[2]==1: wsheet.cell(Row,column_index_from_string("F")).value=2

    i=findNum(doc.paragraphs,i+1,"企业种植基地")
    a=[1]*5
    for j in range(len(doc.paragraphs[i].text)):
        if doc.paragraphs[i].text[j]=='□':
            if j<=len(doc.paragraphs[i].text)-2 and doc.paragraphs[i].text[j+1]=='√':continue;
            str1=doc.paragraphs[i].text[j-4:j]
            if str1=='种植基地':a[1]=0
            elif str1=='作社成员':a[2]=0
            elif str1=='作社领导':a[3]=0
            elif str1=='其他农户':a[4]=0
    x=0
    #print(a)
    for j in range(1,5):
        if a[j]==1:
            x=x*10+j
    if x!=0:wsheet.cell(Row,column_index_from_string("K")).value=x
    
    
    i=findNum(doc.paragraphs,i+1,"重要经济来源")
    a=[1]*3
    for j in range(len(doc.paragraphs[i].text)):
        if doc.paragraphs[i].text[j]=='□':
            if j<=len(doc.paragraphs[i].text)-2 and doc.paragraphs[i].text[j+1]=='√':continue;
            if doc.paragraphs[i].text[j-1]=='否':a[2]=0
            elif doc.paragraphs[i].text[j-1]=='是':a[1]=0
    if a[1]==1: wsheet.cell(Row,column_index_from_string("M")).value=1
    if a[2]==1: wsheet.cell(Row,column_index_from_string("M")).value=2

    i=findNum(doc.paragraphs,i+1,"平均年龄")
    a=[1]*6
    for j in range(len(doc.paragraphs[i].text)):
        if doc.paragraphs[i].text[j]=='□':
            if j<=len(doc.paragraphs[i].text)-2 and doc.paragraphs[i].text[j+1]=='√':
                continue;
            str1=doc.paragraphs[i].text[j-2:j]
            if str1=='30':a[1]=0
            elif str1=='40':a[2]=0
            elif str1=='50':a[3]=0
            elif str1=='60':a[4]=0
            elif str1=='以上':a[5]=0
    for j in range(1,6):
        if a[j]==1:
            wsheet.cell(Row,column_index_from_string("N")).value=j

    i=findNum(doc.paragraphs,i+1,"最高学历")
    a=[1]*6
    for j in range(len(doc.paragraphs[i].text)):
        if doc.paragraphs[i].text[j]=='□':
            if j<=len(doc.paragraphs[i].text)-2 and doc.paragraphs[i].text[j+1]=='√':continue;
            str1=doc.paragraphs[i].text[j-2:j]
            if str1=='以下':a[1]=0
            elif str1=='小学':a[2]=0
            elif str1=='中学':a[3]=0
            elif str1=='本科':a[4]=0
            elif str1=='以上':a[5]=0
    for j in range(1,6):
        if a[j]==1:
            wsheet.cell(Row,column_index_from_string("O")).value=j

    i=findNum(doc.paragraphs,i+1,"加入合作社")
    a=[1]*6
    for j in range(len(doc.paragraphs[i].text)):
        if doc.paragraphs[i].text[j]=='□':
            if j<=len(doc.paragraphs[i].text)-2 and doc.paragraphs[i].text[j+1]=='√':continue;
            str1=doc.paragraphs[i].text[j-1:j]
            if str1=='是':a[1]=0
            elif str1=='否':a[2]=0
    if a[2]==1:
        wsheet.cell(Row,column_index_from_string("P")).value=2
    if a[1]==1:
        wsheet.cell(Row,column_index_from_string("P")).value=1
        for j in range(len(doc.paragraphs[i].text)):
            if doc.paragraphs[i].text[j]=='称':
                str1=doc.paragraphs[i].text[j+1:]
                break;
        for j in range(len(str1)):
            if str1[j]=='否':
                str1=str1[:j]
                break;
        str1=str1.replace(' ','')
        if str1!=[]:
            wsheet.cell(Row,column_index_from_string("Q")).value=str1
    
    i=findNum(doc.paragraphs,i+1,"耕作制度")
    a=[1]*6
    for j in range(len(doc.paragraphs[i].text)):
        if doc.paragraphs[i].text[j]=='□':
            if j<=len(doc.paragraphs[i].text)-2 and doc.paragraphs[i].text[j+1]=='√':continue;
            str1=doc.paragraphs[i].text[j-2:j]
            if str1=='一熟':a[1]=0
            elif str1=='两熟':a[2]=0
    if a[2]==1:
        wsheet.cell(Row,column_index_from_string("T")).value=2
    if a[1]==1:
        wsheet.cell(Row,column_index_from_string("T")).value=1
    if a[1]+a[2]==0:
        for j in range(len(doc.paragraphs[i].text)):
            if doc.paragraphs[i].text[j]=='）':
                str1=doc.paragraphs[i].text[j+1:]
                break;
        str1=str1.replace(' ','')
        if str1!=[]:
            wsheet.cell(Row,column_index_from_string("U")).value=str1

            
    tbs = doc.tables
    tableNum=0
    index0=column_index_from_string("V")

    tb = tbs[tableNum]
    a=[]
    str1=''
    for j in range(1,len(tb.rows)-1):
        for k in range(len(tb.rows[j].cells)):
            temp=tb.rows[j].cells[k].text
            if  temp!= str1:
                if temp=='' and a==[]: continue;
                a.append(temp)
                str1= temp
                if len(a)==3:
                    wsheet.cell(Row,index0).value=a[0]
                    wsheet.cell(Row,index0+1).value=a[1]
                    wsheet.cell(Row,index0+2).value=a[2]
                    a=[]
                    str1=''
                    index0=index0+3
                #print(str(i+1),str(j+1),str(k+1),tb.rows[j].cells[k].text)
                #str1=tb.rows[j].cells[k].text
    j=len(tb.rows)-1;
    for k in range(len(tb.rows[j].cells)):
        temp=tb.rows[j].cells[k].text
        #print(temp)
        if '园总面积' in temp:
            #print(type(temp.find('：')),temp.find('亩'))
            temp2=temp[temp.find('：')+1:temp.find('亩')].replace(' ','')
            if temp2!='':
                wsheet.cell(Row,column_index_from_string("AT")).value=Digit(temp2)
            break;

    tb = tbs[tableNum+1]            
    i=0;

    [i,j]=tableFindNum(tb,i,'建立时间')
    temp=tb.rows[i].cells[j].text;
    temp2=temp[temp.find('：')+1:].replace(' ','')
    wsheet.cell(Row,column_index_from_string("AU")).value=Digit(temp2)
    #print(tb.rows[i].cells[j].text)

    [i,j]=tableFindNum(tb,i,'可使用年限')
    temp=tb.rows[i].cells[j].text;
    temp2=temp[temp.find('：')+1:].replace(' ','')
    wsheet.cell(Row,column_index_from_string("AV")).value=Digit(temp2)

    [i,j]=tableFindNum(tb,i,'租用土地建葡萄园')
    temp=tb.rows[i].cells[j].text;
    temp2=temp[temp.find('：')+1:temp.find('元/亩')].replace(' ','')
    wsheet.cell(Row,column_index_from_string("AW")).value=Digit(temp2)

    [i,j]=tableFindNum(tb,i,'自家土地建葡萄园')
    temp=tb.rows[i].cells[j].text;
    temp2=temp[temp.find('：')+1:temp.find('元/亩')].replace(' ','')
    wsheet.cell(Row,column_index_from_string("AX")).value=Digit(temp2)

    [i,j]=tableFindNum(tb,i,'建园物质成本')
    for x in range(j+1,len(tb.rows[i].cells)):
        if tb.rows[i].cells[x].text!= tb.rows[i].cells[j].text:
            j=x;break;        
    temp=tb.rows[i].cells[j].text;
    #if temp!='' and temp[0]=='7':temp=tb.rows[i].cells[j].text[1:];
    if '左家综合试验站'in word_file and '30个非示范园' in word_file or "南京站" in word_file :
        temp=temp[1:];
    temp2=temp.replace(' ','')
    #print(tb.rows[i].cells[j].text,temp2,'建园物质成本')
    wsheet.cell(Row,column_index_from_string("AY")).value=Digit(temp2)

    [i,j]=tableFindNum(tb,i,'建园人工成本')
    for x in range(j+1,len(tb.rows[i].cells)):
        if tb.rows[i].cells[x].text!= tb.rows[i].cells[j].text:
            j=x;break;        
    temp=tb.rows[i].cells[j].text;
    temp2=temp.replace(' ','')
    wsheet.cell(Row,column_index_from_string("AZ")).value=Digit(temp2)

    [i,j]=tableFindNum(tb,i,'新增维护费用')
    temp=tb.rows[i].cells[j].text;
    temp2=temp[temp.find('：')+1:temp.find('元/亩')].replace(' ','')
    wsheet.cell(Row,column_index_from_string("BA")).value=Digit(temp2)

    [i,j]=tableFindNum(tb,i,'花费资金')
    flag=0
    str1=''
    for x in range(i,len(tb.rows)-1):
        if x==i: j0=j
        else: j0=0
        for y in range(j0,len(tb.rows[x].cells)):
            temp=tb.rows[x].cells[y].text
            if '材料名称' not in temp and '花费资金' in str1 and '花费资金' not in temp:
                i=x; j=y;flag=1;
                break;
            str1=temp
        if flag==1:break;



    tb = tbs[tableNum+2]            
    i=0;
    Count=0
    a=[]
    str1=''
    [i,j]=tableFindNum(tb,i,'水')
    index0=column_index_from_string("CC")
    for x in range(i,len(tb.rows)):
        if x==i: j0=j
        else: j0=0
        for y in range(j0,len(tb.rows[x].cells)):
            temp=tb.rows[x].cells[y].text
            #print(temp)
            if '投入花费资金' in temp:break;  
            if  str1!=temp:
                a.append(temp)
                if len(a)==3:
                    if is_AllChinese(a[0]) and is_AllChinese(a[1]) and is_AllChinese(a[2]):
                        a=a[1:]
                    elif Count<8:
                        if Cost(a[0])>=0 and Cost(a[1])<0 and Cost(a[2])<0:
                            Count=Count+1
                            wsheet.cell(Row,index0+Cost(a[0])*3).value=Digit(a[0])
                            wsheet.cell(Row,index0+Cost(a[0])*3+1).value=Digit(a[1])
                            wsheet.cell(Row,index0+Cost(a[0])*3+2).value=Digit(a[2])
                            a=[]
                        else: a=a[1:]
                    else:
                        Count=Count+1
                        wsheet.cell(Row,index0+Count*3-3).value=Digit(a[0])
                        wsheet.cell(Row,index0+Count*3-3+1).value=Digit(a[1])
                        wsheet.cell(Row,index0+Count*3-3+2).value=Digit(a[2])
                        a=[];str1='';

            str1=temp;
        if '投入花费资金' in temp:
            [i,j]=tableFindNum(tb,x,'投入花费资金')
            #print(i,j,tb.rows[i].cells[j].text)
            break;

    temp=tb.rows[i].cells[j].text;
    temp2=temp[temp.find('：')+1:temp.find('（')].replace(' ','')
    wsheet.cell(Row,column_index_from_string("DS")).value=Digit(temp2)
    #num_Chinese(word)

    [i,j]=tableFindNum(tb,i,'机械寿')
    #print(i,j,tb.rows[i].cells[j].text)
    flag=0
    for x in range(i,len(tb.rows)):
        j0=0    
        if x==i: j0=j
        for y in range(j0,len(tb.rows[x].cells)):
            temp=tb.rows[x].cells[y].text
            #print(x,y,temp)
            if '机械寿' not in temp:
                i=x; j=y;flag=1
                break; 
        if flag==1:
            break; 
    a=[]
    str1=''
    index0=column_index_from_string("DT")
    #print(i,j,tb.rows[i].cells[j].text,'!')
    for x in range(i,len(tb.rows)):
        j0=0
        if x==i: j0=j
        for y in range(j0,len(tb.rows[x].cells)):
            temp=tb.rows[x].cells[y].text
            #print(temp)
            if '农机具作业' in temp:break;  
            if  str1!=temp:
                a.append(temp)
                str1=temp;
                #print(a)
                if len(a)==6:
                    z0=0;
                    for z in range(5,0,-1):
                        if is_AllChinese(a[z])==False:
                            z0=z;
                            break;
                    for z in range(z0+1):
                        p=z;
                        if z>=3:
                            p=p+1
                        wsheet.cell(Row,index0+p).value=Digit(a[z]);#print(p,z,'^',a[z])
                    
                    index0=index0+7
                    a=a[z0+1:]
                    #str1='';
        if '农机具作业' in temp:break;
    #print(i,j,tb.rows[i].cells[j].text)

    
    index0=column_index_from_string("EO")
    [i,j]=tableFindNum(tb,i,'农机具作业')
    temp=tb.rows[i].cells[j].text;
    temp2=temp[temp.find('：')+1:temp.find('（')].replace(' ','')
    wsheet.cell(Row,index0).value=Digit(temp2)
    
    [i,j]=tableFindNum(tb,i,'备注')
    str1=''
    a=[]
    index0=column_index_from_string("EP")
    for x in range(i,len(tb.rows)):
        j0=0
        if x==i: j0=j
        for y in range(j0,len(tb.rows[x].cells)):
            temp=tb.rows[x].cells[y].text.replace(' ','')
            if '每年自家投入的人工' in temp:break;
            if '生产全过程' in str1 and '生产全过程' not in temp:
                i=x;
                j=y
                a.append(temp)
                break;  
            str1=temp;
        if a!=[]:
            break;
    
    a.append('')
    for x in range(i,len(tb.rows)):
        j0=0
        if x==i: j0=j
        for y in range(j0,len(tb.rows[x].cells)):            
            temp=tb.rows[x].cells[y].text.replace(' ','')
            if '每年自家投入的人工' in temp:break;
            if temp!=a[0]:
                a.append('')
                a[1]=temp;
                i=x; j=y
                break;  
            str1=temp;
        if temp!=a[0]:
            break;
    wsheet.cell(Row,index0).value=Digit(a[0]);
    wsheet.cell(Row,index0+1).value=Digit(a[1]);
    
    [i,j]=tableFindNum(tb,i,'备注')
    str1=''
    a=[]
    index0=column_index_from_string("ER")
    for x in range(i,len(tb.rows)):
        j0=0
        if x==i: j0=j
        for y in range(j0,len(tb.rows[x].cells)):
            temp=tb.rows[x].cells[y].text.replace(' ','')
            if '生产全过程' in str1 and '生产全过程' not in temp:
                i=x;
                j=y
                a.append(temp)
                break;  
            str1=temp;
        if a!=[]:
            break;
    
    a.append('')
    for x in range(i,len(tb.rows)):
        j0=0
        if x==i: j0=j
        for y in range(j0,len(tb.rows[x].cells)):            
            temp=tb.rows[x].cells[y].text.replace(' ','')
            if temp!=a[0]:
                a.append('')
                a[1]=temp;
                i=x; j=y
                break;  
            str1=temp;
        if temp!=a[0]:
            break;
    wsheet.cell(Row,index0).value=Digit(a[0]);
    wsheet.cell(Row,index0+1).value=Digit(a[1]);
    
    index0=column_index_from_string("FP")
    i=findNum(doc.paragraphs,len(doc.paragraphs)-3,"调查员")
    #[i,j]=tableFindNum(tb,i,'')
    if i>1:        
        temp=doc.paragraphs[i].text
        if "姓名" in temp:
            temp2=temp[temp.find('名：')+2:temp.find('调研时')].replace(' ','')
            wsheet.cell(Row,index0).value=Digit(temp2)
        #print(i,j,temp2)
        
        index0=index0+1
        if "时间" in temp:
            temp2=temp[temp.find('间：')+2:temp.find('调研地')].replace(' ','')
            wsheet.cell(Row,index0).value=Digit(temp2)
        
        index0=index0+1
        if "地点" in temp:
            temp2=temp[temp.find('点：')+2:].replace(' ','')
            wsheet.cell(Row,index0).value=Digit(temp2)
    
    
    w = win32com.client.Dispatch('Word.Application')
    # 或者使用下面的方法，使用启动独立的进程：
    # w = win32com.client.DispatchEx('Word.Application')

    # 后台运行，显示程序界面，不警告
    w.Visible = 0 #0：后台运行，不显示文档。这个至少在调试阶段建议打开，否则如果等待时间长的话，它至少给你耐心。。。
    w.DisplayAlerts = 0 #0：不警告

    # 打开新的文件
    worddoc = w.Documents.Open(word_file)
    wtb=worddoc.Tables[1]
    
    try:
        for i in range(len(wtb.Rows)-1,len(wtb.Rows)-7,-1):
            if '材料名称' in wtb.Rows[i].Cells[0].Range.Text:
                x=i+1;break;
        index0=column_index_from_string("BB")
        index1=column_index_from_string("BN")
        for i in range(x,len(wtb.Rows)):
            str1=wtb.Rows[i].Cells[0].Range.Text[:-1].replace(' ','')
            #print(i,str1)
            if '合计' in str1 or '共计' in str1 or len(wtb.Rows[i].Cells)<=3:break;
            if str1!='' and '合计' not in str1:
                wsheet.cell(Row,index0).value=str1;
                wsheet.cell(Row,index0+1).value=Digit(wtb.Rows[i].Cells[1].Range.Text[:-1].replace(' ',''));
                wsheet.cell(Row,index0+2).value=Digit(wtb.Rows[i].Cells[2].Range.Text[:-1].replace(' ',''));
                index0=index0+3
            str1=wtb.Rows[i].Cells[3].Range.Text[:-1].replace(' ','')
            #print(i,str1)
            if '合计' in str1 or '共计' in str1 or len(wtb.Rows[i].Cells)<6:break;
            if str1!='' :
                wsheet.cell(Row,index1).value=str1;
                wsheet.cell(Row,index1+1).value=Digit(wtb.Rows[i].Cells[4].Range.Text[:-1].replace(' ',''));
                wsheet.cell(Row,index1+2).value=Digit(wtb.Rows[i].Cells[5].Range.Text[:-1].replace(' ',''));
                index1=index1+3
    except :
        pass;
    
    


#四、葡萄园的收益情况             
    try:
        index0=column_index_from_string("ET")
        wtb=worddoc.Tables[3]
        for i in range(1,len(wtb.Rows)-1):
            if i>=5:break;
            #for j in range(0,4):
            wsheet.cell(Row,index0).value=Digit(wtb.Rows[i].Cells[0].Range.Text[:-1]);
            wsheet.cell(Row,index0+1).value=Digit(wtb.Rows[i].Cells[1].Range.Text[:-1]);
            wsheet.cell(Row,index0+2).value=Digit(wtb.Rows[i].Cells[2].Range.Text[:-1]);
            if i==4:break;
            wsheet.cell(Row,index0+12).value=Digit(wtb.Rows[i].Cells[3].Range.Text[:-1]);
            wsheet.cell(Row,index0+12+1).value=Digit(wtb.Rows[i].Cells[4].Range.Text[:-1]);
            wsheet.cell(Row,index0+12+2).value=Digit(wtb.Rows[i].Cells[5].Range.Text[:-1]);
            index0=index0+3

        index0=column_index_from_string("FO")
        temp=wtb.Rows[len(wtb.Rows)-1].Cells[0].Range.Text[:-1]
        #print(temp)
        temp=temp[temp.find('是')+1:temp.find('元')].replace(' ','')
        wsheet.cell(Row,index0).value=Digit(temp);
    except :
        pass;
    
    
    
    
    
    worddoc.Close()   #关闭文档
    w.Quit() #关闭word程序
    wbook.save(xlsx_file)
    
def ExtractRecursion(Path,xlsx_file=''):#传入文件夹和模板的地址，将该文件夹的所有问卷文件的信息填入模板，再递归地处理所有子文件夹
    global CompletedNum,Dict,Template_xlsx
    if xlsx_file=='':
        xlsx_file=Template_xlsx;
    file_folder=next(os.walk(Path));
    #if len(file_folder[1])>0:#file_folder[1]是子文件夹
    for i in file_folder[1]:
        ExtractRecursion(os.path.join(Path,i));
    Count=0
    Fail=[]
    for i in file_folder[2]:
        path=os.path.join(Path,i);
        if os.path.splitext(path)[1] == ".docx":
            Count=Count+1
            try:
                Extract(path,xlsx_file);
            except Exception as e:
                Fail.append(path)
                print(Count,"失败!",e)
                traceback.print_exc()#异常信息，打断程序运行
                #print(sys.exc_info()) 异常信息，不打断程序运行
            else:
                print(Count,"成功!")        
            print()   
    return 0

def Automatic_entry(File1,File2,File3):
    global CompletedNum,Dict,Template_xlsx
    #File1 文件夹所在文件夹
    #File2 模板文件
    #File3 汇总表保存到的文件夹
    File3=os.path.join(File3, '汇总表.xlsx')
    File3=File3.replace('\\','//')

    wbook=openpyxl.load_workbook(File2.replace('\\','//'))
    wbook.save(File3)
    Template_xlsx=File3
    General_file_folder=File1.replace('\\','//')
    wbook=openpyxl.load_workbook(Template_xlsx)
    Name=wbook.sheetnames
    CompletedNum=[0]*len(Name) #已经填写数，如CompletedNum[6]=9表示编号为9的省份已经填写了9份问卷
    Dict = {}
    for i in range(len(Name)):
        Dict[Name[i]]=i #省份序号，如Dict['广西']为7

    ExtractRecursion(General_file_folder)

import sys
from PyQt5.QtWidgets import QApplication, QWidget, QPushButton, QAction, QMessageBox
from PyQt5.QtGui import QIcon
from PyQt5.QtCore import pyqtSlot
from PyQt5.Qt import QLineEdit

from PyQt5.QtCore import Qt
from PyQt5.QtWidgets import QLineEdit
from PyQt5.QtWidgets import QTextEdit
from PyQt5.QtWidgets import QSizePolicy
from PyQt5.QtWidgets import QMainWindow
from PyQt5.QtWidgets import QPushButton
from PyQt5.QtWidgets import QGridLayout
from PyQt5.QtWidgets import QApplication
from PyQt5.QtWidgets import QFrame
from PyQt5.QtWidgets import QLabel
from PyQt5.QtWidgets import QWidget

class App(QWidget):

    def __init__(self):
        super().__init__()
        self.title = '葡萄成本收益调查问卷数据自动录入系统V1.0'
        self.left = 300
        self.top = 400
        self.width = 840
        self.height = 440
        self.initUI()

    def initUI(self):
        self.setWindowTitle(self.title)
        self.setGeometry(self.left, self.top, self.width, self.height)



        x1=60
        y1=60
        self.textbox = QLineEdit(self)
        self.textbox.move(x1+100,y1)
        self.textbox.resize(280, 30)
        self.textbox.setText("")
        label = QLabel(self)
        label.resize(100, 30)
        label.move(x1, y1+5)
        label.setText("问卷所在文件夹：")
        label.setAlignment(Qt.AlignRight)

        x2 = x1
        y2 = y1+100
        self.textbox2 = QLineEdit(self)
        self.textbox2.move(x2+100, y2)
        self.textbox2.resize(280, 30)
        self.textbox2.setText("")
        label2 = QLabel(self)
        label2.resize(100, 30)
        label2.move(x2, y2+5)
        label2.setText("模板文件：")
        label2.setAlignment(Qt.AlignRight)

        x3 = x1
        y3 = y2 + 100
        self.textbox3 = QLineEdit(self)
        self.textbox3.move(x3 + 100, y3)
        self.textbox3.resize(280, 30)
        self.textbox3.setText("")
        label3 = QLabel(self)
        label3.resize(100, 30)
        label3.move(x3, y3 + 5)
        label3.setText("汇总表保存到：")
        label3.setAlignment(Qt.AlignRight)

        x4 = 500
        y4 = 60
        label4 = QLabel(self)
        label4.resize(260, 180)
        label4.move(x4, y4 + 5)
        label4.setFrameStyle(QFrame.Panel | QFrame.Sunken)
        str1="软件说明：\n本软件可以读取指定文件夹中所有.docx后缀的问卷文件，自动识别其中相关信息并填入指定的Excel文件中，并另存到指定文件夹中。\n"
        str2 = "\n使用步骤：\n在三个文本框中依次输入存放Word问卷的文件夹地址、Excel模板文件的地址和保存输出汇总表的文件夹地址，"                "再点击“确定”，打开汇总表文件夹即可查看输出汇总结果。"
        label4.setWordWrap(True)
        label4.setText(str1+str2)
        label4.setAlignment(Qt.AlignTop | Qt.AlignLeft)


        # Create a button in the window
        self.button = QPushButton('确定', self)
        self.button.move(250, 350)

        # connect button to function on_click
        self.button.clicked.connect(self.on_click)
        #self.center()
        self.show()

    @pyqtSlot()
    def on_click(self):
        textboxValue1 = self.textbox.text()
        textboxValue2 = self.textbox2.text()
        textboxValue3 = self.textbox3.text()
        print(textboxValue1)
        print(textboxValue2)
        print(textboxValue3)
        Automatic_entry(textboxValue1, textboxValue2, textboxValue3)
        #print(type(textboxValue))
        #self.textbox.setText(textboxValue+"!")
        """打印完毕之后清空文本框"""
        #self.textbox.setText('')


if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = App()
    app.exit(app.exec_())
